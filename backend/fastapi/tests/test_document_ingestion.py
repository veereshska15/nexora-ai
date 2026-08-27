import io
import pytest
from fastapi.testclient import TestClient
import docx

from core.config import settings
from documents.document_service import document_service
from documents.parsers.txt_parser import TxtParser
from documents.parsers.csv_parser import CsvParser
from documents.parsers.pdf_parser import PdfParser
from documents.parsers.docx_parser import DocxParser

# Sample valid PDF with extracted text content
VALID_TEST_PDF = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 53 >>
stream
BT
/F1 12 Tf
72 712 Td
(NEXORA AI Multilingual Document Test) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000244 00000 n 
0000000347 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
424
%%EOF"""


def generate_test_docx(paragraphs: list[str], table_rows: list[list[str]] = None) -> bytes:
    """Helper to generate a valid in-memory DOCX document."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row in enumerate(table_rows):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==============================================================================
# 1. PARSER LEVEL UNIT TESTS
# ==============================================================================

def test_txt_extraction():
    content = "Welcome to NEXORA AI.\nMultilingual AI Platform for Enterprise."
    file_bytes = content.encode("utf-8")
    parser = TxtParser()
    res = parser.extract(file_bytes, "intro.txt")

    assert res.extraction_success is True
    assert res.file_type == "txt"
    assert res.filename == "intro.txt"
    assert res.file_size == len(file_bytes)
    assert "NEXORA AI" in res.extracted_text
    assert res.character_count == len(res.extracted_text)
    assert res.page_count is None
    assert res.processing_time_ms >= 0.0


def test_csv_extraction():
    content = "id,name,role\n1,Veeresh,Architect\n2,Nexora,AI Assistant"
    file_bytes = content.encode("utf-8")
    parser = CsvParser()
    res = parser.extract(file_bytes, "users.csv")

    assert res.extraction_success is True
    assert res.file_type == "csv"
    assert "id | name | role" in res.extracted_text
    assert "1 | Veeresh | Architect" in res.extracted_text
    assert res.character_count > 0


def test_pdf_extraction():
    parser = PdfParser()
    res = parser.extract(VALID_TEST_PDF, "sample.pdf")

    assert res.extraction_success is True
    assert res.file_type == "pdf"
    assert res.page_count == 1
    assert "NEXORA AI Multilingual Document Test" in res.extracted_text
    assert res.character_count > 0


def test_docx_extraction():
    docx_bytes = generate_test_docx(
        paragraphs=["NEXORA AI Architecture Overview", "Deep Neural Networks and Vector DB"],
        table_rows=[["Service", "Port"], ["FastAPI", "8000"], ["Qdrant", "6333"]],
    )
    parser = DocxParser()
    res = parser.extract(docx_bytes, "architecture.docx")

    assert res.extraction_success is True
    assert res.file_type == "docx"
    assert "NEXORA AI Architecture Overview" in res.extracted_text
    assert "FastAPI | 8000" in res.extracted_text
    assert res.character_count > 0


# ==============================================================================
# 2. VALIDATION & SECURITY TESTS
# ==============================================================================

def test_empty_file_rejection():
    res = document_service.extract_from_bytes(b"", "empty.txt")
    assert res.extraction_success is False
    assert "empty" in res.error_message.lower()


def test_unsupported_extension_rejection():
    res = document_service.extract_from_bytes(b"malicious payload", "script.exe")
    assert res.extraction_success is False
    assert "Unsupported file format" in res.error_message


def test_file_size_validation():
    oversized_bytes = b"X" * (settings.MAX_UPLOAD_SIZE_BYTES + 1024)
    res = document_service.extract_from_bytes(oversized_bytes, "large.txt")
    assert res.extraction_success is False
    assert "exceeds maximum allowed upload limit" in res.error_message


def test_metadata_generation():
    text = "Detailed system metadata test document."
    file_bytes = text.encode("utf-8")
    res = document_service.extract_from_bytes(file_bytes, "meta_test.txt")

    assert res.filename == "meta_test.txt"
    assert res.file_type == "txt"
    assert res.file_size == len(file_bytes)
    assert res.character_count == len(text)
    assert res.processing_time_ms >= 0.0
    assert res.error_message is None


# ==============================================================================
# 3. UNICODE & KANNADA TEXT EXTRACTION
# ==============================================================================

def test_unicode_kannada_extraction_txt():
    kannada_text = "ನಮಸ್ಕಾರ, NEXORA AI ಗೆ ಸುಸ್ವಾಗತ. ಇದು ಕನ್ನಡ ಪಠ್ಯ ಸಂಸ್ಕರಣಾ ವ್ಯವಸ್ಥೆ."
    file_bytes = kannada_text.encode("utf-8")
    res = document_service.extract_from_bytes(file_bytes, "kannada_sample.txt")

    assert res.extraction_success is True
    assert "ನಮಸ್ಕಾರ" in res.extracted_text
    assert "ಕನ್ನಡ" in res.extracted_text


def test_unicode_kannada_extraction_docx():
    kannada_paras = [
        "ಕನ್ನಡ ಭಾಷೆ ಮತ್ತು ಕೃತಕ ಬುದ್ಧಿಮತ್ತೆ",
        "ನರಮಂಡಲ ಜಾಲಗಳ ವಿಶ್ಲೇಷಣೆ",
    ]
    docx_bytes = generate_test_docx(kannada_paras)
    res = document_service.extract_from_bytes(docx_bytes, "kannada_doc.docx")

    assert res.extraction_success is True
    assert "ಕೃತಕ ಬುದ್ಧಿಮತ್ತೆ" in res.extracted_text


# ==============================================================================
# 4. FASTAPI REST API ENDPOINT INTEGRATION TESTS
# ==============================================================================

def test_api_extract_endpoint_txt(client: TestClient):
    file_content = b"FastAPI Multipart Document Extraction Test"
    files = {"file": ("test.txt", file_content, "text/plain")}

    response = client.post("/api/v1/documents/extract", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "test.txt"
    assert data["file_type"] == "txt"
    assert data["extracted_text"] == "FastAPI Multipart Document Extraction Test"
    assert data["character_count"] == len("FastAPI Multipart Document Extraction Test")
    assert data["extraction_success"] is True
    assert data["processing_time_ms"] >= 0.0


def test_api_extract_endpoint_csv(client: TestClient):
    file_content = b"language,script,is_indic\nKannada,Kannada,true\nHindi,Devanagari,true"
    files = {"file": ("languages.csv", file_content, "text/csv")}

    response = client.post("/api/v1/documents/extract", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "languages.csv"
    assert data["file_type"] == "csv"
    assert "Kannada | Kannada | true" in data["extracted_text"]


def test_api_extract_endpoint_pdf(client: TestClient):
    files = {"file": ("report.pdf", VALID_TEST_PDF, "application/pdf")}

    response = client.post("/api/v1/documents/extract", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "report.pdf"
    assert data["file_type"] == "pdf"
    assert data["page_count"] == 1
    assert "NEXORA AI Multilingual Document Test" in data["extracted_text"]


def test_api_extract_endpoint_docx(client: TestClient):
    docx_bytes = generate_test_docx(["DOCX API Endpoint Test", "Structured Word Document"])
    files = {
        "file": (
            "test_doc.docx",
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    response = client.post("/api/v1/documents/extract", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "test_doc.docx"
    assert data["file_type"] == "docx"
    assert "DOCX API Endpoint Test" in data["extracted_text"]


def test_api_extract_empty_file_rejected(client: TestClient):
    files = {"file": ("empty.txt", b"", "text/plain")}

    response = client.post("/api/v1/documents/extract", files=files)
    assert response.status_code == 400
    assert "empty" in response.json()["detail"].lower()


def test_api_extract_unsupported_format_rejected(client: TestClient):
    files = {"file": ("binary.bin", b"\x00\x01\x02\x03", "application/octet-stream")}

    response = client.post("/api/v1/documents/extract", files=files)
    assert response.status_code == 400
    assert "Unsupported file format" in response.json()["detail"]
