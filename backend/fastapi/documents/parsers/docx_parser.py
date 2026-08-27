import io
from typing import Optional
import docx
from documents.parsers.base_parser import BaseDocumentParser


class DocxParser(BaseDocumentParser):
    """
    Parser for Microsoft Word documents (.docx) using python-docx.
    Extracts text across paragraphs and embedded tables.
    """

    def __init__(self):
        super().__init__(file_type="docx")

    def _parse_content(self, file_bytes: bytes, filename: str) -> tuple[str, Optional[int]]:
        docx_stream = io.BytesIO(file_bytes)
        doc = docx.Document(docx_stream)

        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        extracted = "\n\n".join(text_parts).strip()
        return extracted, None
